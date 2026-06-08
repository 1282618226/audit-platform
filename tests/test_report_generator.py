"""测试 report_generator.py —— CNAS 审计报告生成。

所有测试写入临时目录，不污染项目目录。
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import pytest

from src.orchestrator import ScanMetadata, ScanResult
from src.report_generator import ReportGenerator


# ─── helpers ──────────────────────────────────────────────────────


def _make_result(
    findings: list[dict] | None = None,
    mode: str = "online",
    run_id: str = "test-run-001",
    duration: float = 42.5,
    metadata: ScanMetadata | None = None,
    llm_reviewed: list[dict] | None = None,
    warnings: list[str] | None = None,
) -> ScanResult:
    """快速构造 ScanResult。"""
    result = ScanResult()
    result.run_id = run_id
    result.mode = mode
    result.duration_seconds = duration
    result.findings = findings or []
    result.llm_reviewed = llm_reviewed or []
    result.warnings = warnings or []
    if metadata:
        result.metadata = metadata
    else:
        m = ScanMetadata()
        m.languages_detected = ["Java"]
        m.total_files = 10
        m.java_files = 10
        m.build_system = "maven"
        m.compile_ready = True
        result.metadata = m
    return result


def _sample_findings() -> list[dict]:
    """构造一组示例发现。"""
    return [
        {
            "clause": "6.2.3.4",
            "standard": "GB/T 34944-2017",
            "vuln_name": "SQL注入",
            "category": "数据处理",
            "file_path": "src/LoginServlet.java",
            "line_start": 42,
            "line_end": 48,
            "source_tool": "semgrep+codeql",
            "auto_confidence": 0.95,
            "code_snippet": 'String query = "SELECT * FROM users WHERE id=\'" + uid + "\'";',
            "tool_raw_output": {"check_id": "java.sql.injection"},
            "severity": "高",
        },
        {
            "clause": "6.2.6.3",
            "standard": "GB/T 34944-2017",
            "vuln_name": "口令硬编码",
            "category": "安全功能",
            "file_path": "src/Config.java",
            "line_start": 10,
            "line_end": 12,
            "source_tool": "semgrep",
            "auto_confidence": 0.95,
            "code_snippet": 'if ("admin123".equals(password))',
            "tool_raw_output": {"check_id": "hardcoded.password"},
            "severity": "中",
        },
        {
            "clause": "6.2.9.1",
            "standard": "GB/T 34944-2017",
            "vuln_name": "点击劫持",
            "category": "用户界面错误",
            "file_path": "src/PageServlet.java",
            "line_start": 15,
            "line_end": 20,
            "source_tool": "semgrep",
            "auto_confidence": 0.55,  # 低置信度 → 疑似
            "code_snippet": "// missing X-Frame-Options",
            "tool_raw_output": {"check_id": "clickjacking"},
            "severity": "低",
        },
    ]


# ─── fixtures ─────────────────────────────────────────────────────


@pytest.fixture
def gen() -> ReportGenerator:
    return ReportGenerator(kb=None)


@pytest.fixture
def gen_with_kb(kb_path: Path) -> ReportGenerator:
    from src.knowledge_base import KnowledgeBase
    kb = KnowledgeBase(kb_path)
    return ReportGenerator(kb=kb)


@pytest.fixture
def output_dir() -> Path:
    return Path(tempfile.mkdtemp())


# ─── JSON 报告 ────────────────────────────────────────────────────


class TestJsonReport:
    """测试 JSON 格式报告生成。"""

    def test_generate_json_structure(self, gen: ReportGenerator, output_dir: Path) -> None:
        """JSON 报告应包含所有顶层字段。"""
        result = _make_result(findings=_sample_findings())
        gen.generate(result, output_dir)

        with open(output_dir / "report.json", encoding="utf-8") as f:
            data = json.load(f)

        assert "meta" in data
        assert "audit_scope" in data
        assert "summary" in data
        assert "findings" in data
        assert "llm_reviewed" in data
        assert data["meta"]["run_id"] == "test-run-001"
        assert data["meta"]["mode"] == "online"

    def test_json_empty_findings(self, gen: ReportGenerator, output_dir: Path) -> None:
        """无发现时 JSON 报告应正常生成。"""
        result = _make_result(findings=[])
        gen.generate(result, output_dir)

        with open(output_dir / "report.json", encoding="utf-8") as f:
            data = json.load(f)

        assert data["findings"] == []
        assert data["summary"]["total"] == 0

    def test_json_summary_counts(self, gen: ReportGenerator, output_dir: Path) -> None:
        """汇总统计应正确。"""
        result = _make_result(findings=_sample_findings())
        gen.generate(result, output_dir)

        with open(output_dir / "report.json", encoding="utf-8") as f:
            data = json.load(f)

        summary = data["summary"]
        assert summary["total"] == 3
        assert summary["confirmed"] == 2  # confidence >= 0.7
        assert summary["suspects"] == 1   # confidence 0.4-0.7
        assert summary["by_severity"]["高"] == 1
        assert summary["by_severity"]["中"] == 1
        assert summary["by_severity"]["低"] == 1


# ─── Markdown 报告 ────────────────────────────────────────────────


class TestMarkdownReport:
    """测试 Markdown 格式报告生成。"""

    def test_generate_markdown_structure(self, gen: ReportGenerator, output_dir: Path) -> None:
        """Markdown 报告应包含所有章节标题。"""
        result = _make_result(findings=_sample_findings())
        gen.generate(result, output_dir)

        content = (output_dir / "report.md").read_text(encoding="utf-8")

        assert "# CNAS 源代码安全审计报告" in content
        assert "一、审计概要" in content
        assert "二、覆盖矩阵摘要" in content
        assert "三、漏洞发现汇总" in content
        assert "四、漏洞详情" in content
        assert "五、疑似发现" in content
        assert "六、盲区与局限性" in content
        assert "七、审计元数据" in content

    def test_markdown_contains_online_indicator(self, gen: ReportGenerator, output_dir: Path) -> None:
        """在线模式应提及 LLM 增强。"""
        result = _make_result(findings=[], mode="online")
        gen.generate(result, output_dir)

        content = (output_dir / "report.md").read_text(encoding="utf-8")
        assert "在线模式" in content

    def test_markdown_contains_offline_warning(self, gen: ReportGenerator, output_dir: Path) -> None:
        """离线模式应包含 LLM 限制提示。"""
        result = _make_result(findings=[], mode="offline")
        gen.generate(result, output_dir)

        content = (output_dir / "report.md").read_text(encoding="utf-8")
        assert "离线模式" in content
        # 离线盲区说明
        assert "LLM 增强未执行" in content or "离线" in content

    def test_markdown_empty_findings(self, gen: ReportGenerator, output_dir: Path) -> None:
        """无发现时应显示友好提示。"""
        result = _make_result(findings=[])
        gen.generate(result, output_dir)

        content = (output_dir / "report.md").read_text(encoding="utf-8")
        assert "未发现任何漏洞" in content

    def test_markdown_shows_finding_details(self, gen: ReportGenerator, output_dir: Path) -> None:
        """确认漏洞应显示详细信息。"""
        result = _make_result(findings=_sample_findings())
        gen.generate(result, output_dir)

        content = (output_dir / "report.md").read_text(encoding="utf-8")
        assert "SQL注入" in content
        assert "src/LoginServlet.java" in content
        assert "42" in content

    def test_markdown_shows_suspects(self, gen: ReportGenerator, output_dir: Path) -> None:
        """低置信度的疑似发现应被列出。"""
        result = _make_result(findings=_sample_findings())
        gen.generate(result, output_dir)

        content = (output_dir / "report.md").read_text(encoding="utf-8")
        assert "点击劫持" in content

    def test_markdown_llm_status_shown(self, gen: ReportGenerator, output_dir: Path) -> None:
        """LLM 状态应显示在发现详情中。"""
        findings = [
            {
                "clause": "6.2.3.4", "standard": "GB/T 34944-2017",
                "vuln_name": "SQL注入", "category": "数据处理",
                "file_path": "A.java", "line_start": 1, "line_end": 2,
                "source_tool": "semgrep", "auto_confidence": 0.92,
                "code_snippet": "code", "tool_raw_output": {},
                "severity": "高",
                "llm_status": "confirmed",
                "llm_reasoning": "确实存在注入",
            },
        ]
        result = _make_result(findings=findings)
        gen.generate(result, output_dir)

        content = (output_dir / "report.md").read_text(encoding="utf-8")
        assert "LLM: confirmed" in content

    def test_markdown_with_kb_enriches_details(self, gen_with_kb: ReportGenerator, output_dir: Path) -> None:
        """有 KB 时应补充风险描述和修复建议。"""
        result = _make_result(findings=[{
            "clause": "6.2.3.4", "standard": "GB/T 34944-2017",
            "vuln_name": "SQL注入", "category": "数据处理",
            "file_path": "A.java", "line_start": 1, "line_end": 2,
            "source_tool": "semgrep", "auto_confidence": 0.92,
            "code_snippet": "code", "tool_raw_output": {},
            "severity": "高",
        }])
        gen_with_kb.generate(result, output_dir)

        content = (output_dir / "report.md").read_text(encoding="utf-8")
        assert "风险描述" in content


# ─── HTML 报告 ────────────────────────────────────────────────────


class TestHtmlReport:
    """测试 HTML 格式报告生成。"""

    def test_generate_html(self, gen: ReportGenerator, output_dir: Path) -> None:
        """HTML 报告应生成且包含基本结构。"""
        gen_with_html = ReportGenerator(kb=None, formats=["html"])
        result = _make_result(findings=_sample_findings())
        gen_with_html.generate(result, output_dir)

        path = output_dir / "report.html"
        assert path.exists()
        content = path.read_text(encoding="utf-8")
        assert "<!DOCTYPE html>" in content
        assert "CNAS 源代码安全审计报告" in content
        assert "SQL注入" in content


# ─── 自定义格式 ───────────────────────────────────────────────────


class TestCustomFormats:
    """测试自定义输出格式。"""

    def test_only_json(self, gen: ReportGenerator, output_dir: Path) -> None:
        """仅输出 JSON 格式。"""
        gen_json = ReportGenerator(kb=None, formats=["json"])
        result = _make_result(findings=[])
        files = gen_json.generate(result, output_dir)

        assert "json" in files
        assert "markdown" not in files
        assert (output_dir / "report.json").exists()
        assert not (output_dir / "report.md").exists()

    def test_only_markdown(self, gen: ReportGenerator, output_dir: Path) -> None:
        """仅输出 Markdown 格式。"""
        gen_md = ReportGenerator(kb=None, formats=["markdown"])
        result = _make_result(findings=[])
        files = gen_md.generate(result, output_dir)

        assert "markdown" in files
        assert "json" not in files
        assert (output_dir / "report.md").exists()
        assert not (output_dir / "report.json").exists()


# ─── 边界情况 ────────────────────────────────────────────────────


class TestEdgeCases:
    """测试边界情况。"""

    def test_result_with_warnings(self, gen: ReportGenerator, output_dir: Path) -> None:
        """有警告信息的扫描结果应被记录。"""
        result = _make_result(findings=[], warnings=["CodeQL build failed"])
        gen.generate(result, output_dir)

        with open(output_dir / "report.json", encoding="utf-8") as f:
            data = json.load(f)
        assert len(data["warnings"]) == 1

    def test_llm_reviewed_in_json(self, gen: ReportGenerator, output_dir: Path) -> None:
        """LLM 审查结果应出现在 JSON 中。"""
        llm_reviewed = [
            {
                "finding": {"clause": "6.2.3.4", "file_path": "A.java", "line_start": 1, "line_end": 2},
                "llm_verdict": "confirmed",
                "llm_confidence": 0.92,
                "llm_reasoning": "真实漏洞",
            },
        ]
        result = _make_result(findings=_sample_findings()[:1], llm_reviewed=llm_reviewed)
        gen.generate(result, output_dir)

        with open(output_dir / "report.json", encoding="utf-8") as f:
            data = json.load(f)
        assert len(data["llm_reviewed"]) == 1

    def test_output_dir_created(self, gen: ReportGenerator, tmp_path: Path) -> None:
        """输出目录不存在时应自动创建。"""
        new_dir = tmp_path / "nested" / "output"
        result = _make_result(findings=[])
        gen.generate(result, new_dir)

        assert new_dir.exists()
        assert (new_dir / "report.json").exists()


# ─── DOCX 报告 ────────────────────────────────────────────────────


class TestDocxReport:
    """测试 DOCX 格式报告生成。"""

    def test_generate_docx(self, gen: ReportGenerator, output_dir: Path) -> None:
        """DOCX 报告应生成且包含表格。"""
        gen_with_docx = ReportGenerator(kb=None, formats=["docx"])
        result = _make_result(findings=_sample_findings())
        gen_with_docx.generate(result, output_dir)

        path = output_dir / "report.docx"
        assert path.exists()

        from docx import Document
        doc = Document(str(path))

        # 应包含标题段落
        assert len(doc.paragraphs) >= 2
        assert "安全漏洞" in doc.paragraphs[0].text

        # 应包含表格
        assert len(doc.tables) >= 1
        table = doc.tables[0]

        # 表头 6 列
        assert len(table.rows[0].cells) == 6
        headers = [cell.text for cell in table.rows[0].cells]
        assert "序号" in headers[0]
        assert "安全漏洞名称" in headers[1]
        assert "风险级别" in headers[2]
        assert "截图证明" in headers[5]

        # 数据行 = confirmed 数 (2) + 表头 (1) = 3
        assert len(table.rows) == 3

        # 结束语
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "本结果报告单结束" in all_text

    def test_docx_empty_findings(self, gen: ReportGenerator, output_dir: Path) -> None:
        """无确认漏洞时 DOCX 应只有表头。"""
        gen_with_docx = ReportGenerator(kb=None, formats=["docx"])
        result = _make_result(findings=[])
        gen_with_docx.generate(result, output_dir)

        path = output_dir / "report.docx"
        assert path.exists()

        from docx import Document
        doc = Document(str(path))
        assert len(doc.tables) >= 1
        # 只有表头行
        assert len(doc.tables[0].rows) == 1

    def test_docx_severity_mapping(self, gen: ReportGenerator, output_dir: Path) -> None:
        """风险级别应正确显示。"""
        gen_with_docx = ReportGenerator(kb=None, formats=["docx"])
        findings = [
            {
                "clause": "6.2.3.4", "standard": "GB/T 34944-2017",
                "vuln_name": "SQL注入", "category": "数据处理",
                "file_path": "A.java", "line_start": 1, "line_end": 2,
                "source_tool": "semgrep", "auto_confidence": 0.92,
                "tool_raw_output": {}, "severity": "高",
            },
            {
                "clause": "6.2.9.1", "standard": "GB/T 34944-2017",
                "vuln_name": "点击劫持", "category": "用户界面错误",
                "file_path": "B.java", "line_start": 5, "line_end": 6,
                "source_tool": "semgrep", "auto_confidence": 0.75,
                "tool_raw_output": {}, "severity": "低",
            },
        ]
        result = _make_result(findings=findings)
        gen_with_docx.generate(result, output_dir)

        from docx import Document
        doc = Document(str(output_dir / "report.docx"))
        table = doc.tables[0]

        # 第 1 行数据: 高
        assert table.rows[1].cells[2].text == "高"
        # 第 2 行数据: 低
        assert table.rows[2].cells[2].text == "低"

    def test_docx_default_format_included(self, gen: ReportGenerator, output_dir: Path) -> None:
        """默认格式应包含 docx。"""
        # gen fixture uses defaults (no explicit formats)
        assert "docx" in gen._formats

    def test_docx_location_format(self, gen: ReportGenerator, output_dir: Path) -> None:
        """代码位置格式应为 '文件路径第X行'。"""
        gen_with_docx = ReportGenerator(kb=None, formats=["docx"])
        findings = [
            {
                "clause": "6.2.3.4", "standard": "GB/T 34944-2017",
                "vuln_name": "SQL注入", "category": "数据处理",
                "file_path": "src/LoginServlet.java", "line_start": 42, "line_end": 48,
                "source_tool": "semgrep", "auto_confidence": 0.92,
                "tool_raw_output": {}, "severity": "高",
            },
        ]
        result = _make_result(findings=findings)
        gen_with_docx.generate(result, output_dir)

        from docx import Document
        doc = Document(str(output_dir / "report.docx"))
        table = doc.tables[0]
        location = table.rows[1].cells[4].text
        assert location == "LoginServlet.java第42行"
