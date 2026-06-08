"""报告生成模块。

将 ScanResult 转换为 CNAS 可接受的审计报告格式。

支持格式:
  - Markdown (.md)   — 主报告，评审专家可读
  - JSON (.json)     — 结构化数据，CNAS 系统可消费
  - DOCX (.docx)     — CNAS 能力验证标准格式
  - HTML (.html)     — 可选，浏览器直接查看

报告内容组织（参考设计文档 Section 2.2.6 Phase 5）:
  一、审计概要
  二、覆盖矩阵摘要
  三、漏洞发现汇总（按严重等级 + 按国标大类）
  四、漏洞详情
  五、疑似发现
  六、工具覆盖盲区
  七、审计元数据
"""

from __future__ import annotations

import json
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


class ReportGenerator:
    """CNAS 审计报告生成器。

    用法:
        gen = ReportGenerator(kb=knowledge_base)
        gen.generate(scan_result, output_dir="/workspace/report")
    """

    def __init__(
        self,
        kb: Any | None = None,
        formats: list[str] | None = None,
    ) -> None:
        """初始化报告生成器。

        Args:
            kb: KnowledgeBase 实例，用于补充漏洞信息。
            formats: 输出格式列表，默认 ["json", "markdown"]。
        """
        self._kb = kb
        self._formats = formats or ["json", "markdown", "docx"]

    def generate(self, result: Any, output_dir: str | Path) -> dict[str, str]:
        """生成所有格式的审计报告。

        Args:
            result: ScanResult 实例（来自 Orchestrator.run()）。
            output_dir: 报告输出目录。

        Returns:
            {"format": "file_path"} 字典。
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        files: dict[str, str] = {}

        if "json" in self._formats:
            path = output_dir / "report.json"
            self._generate_json(result, path)
            files["json"] = str(path)

        if "markdown" in self._formats:
            path = output_dir / "report.md"
            self._generate_markdown(result, path)
            files["markdown"] = str(path)

        if "html" in self._formats:
            path = output_dir / "report.html"
            self._generate_html(result, path)
            files["html"] = str(path)

        if "docx" in self._formats:
            path = output_dir / "report.docx"
            self._generate_docx(result, path)
            files["docx"] = str(path)

        return files

    # ─── JSON 报告 ────────────────────────────────────────────────

    def _generate_json(self, result: Any, path: Path) -> None:
        """生成 JSON 格式的结构化报告。"""
        report = {
            "meta": {
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "run_id": getattr(result, "run_id", ""),
                "mode": getattr(result, "mode", "unknown"),
                "duration_seconds": getattr(result, "duration_seconds", 0.0),
            },
            "audit_scope": self._serialize_metadata(result),
            "summary": self._build_summary(result),
            "findings": result.findings if hasattr(result, "findings") else [],
            "llm_reviewed": result.llm_reviewed if hasattr(result, "llm_reviewed") else [],
            "warnings": result.warnings if hasattr(result, "warnings") else [],
        }

        with open(path, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

    # ─── Markdown 报告 ────────────────────────────────────────────

    def _generate_markdown(self, result: Any, path: Path) -> None:
        """生成 Markdown 格式的主报告。"""
        findings = result.findings if hasattr(result, "findings") else []
        metadata = result.metadata if hasattr(result, "metadata") else None

        lines: list[str] = []
        self._md_h1(lines, "CNAS 源代码安全审计报告")
        self._md_section(lines, "一、审计概要", self._render_summary_table(result, metadata))
        self._md_section(lines, "二、覆盖矩阵摘要", self._render_coverage_summary(findings))
        self._md_section(lines, "三、漏洞发现汇总", self._render_findings_summary(findings))
        self._md_section(lines, "四、漏洞详情", self._render_findings_detail(findings))
        self._md_section(lines, "五、疑似发现", self._render_suspect_findings(findings))
        self._md_section(lines, "六、盲区与局限性", self._render_blind_spots(result))
        self._md_section(lines, "七、审计元数据", self._render_metadata(result, metadata))

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    # ─── DOCX 报告（CNAS 格式）─────────────────────────────────────

    def _generate_docx(self, result: Any, path: Path) -> None:
        """生成 CNAS 能力验证标准格式的 DOCX 报告。

        格式: 标题 + 6 列表格 + 结束语。
        只输出确认漏洞（auto_confidence >= 0.7）。
        """
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

        findings = result.findings if hasattr(result, "findings") else []
        confirmed = [f for f in findings if f.get("auto_confidence", 0) >= 0.7]

        doc = Document()

        # ── 默认字体 ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # ── 标题 ──
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run("一、本次发现的源代码安全漏洞如下：")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        doc.add_paragraph()  # 空行

        # ── 表格 ──
        std_label = self._standard_header(result)
        headers = [
            "序号",
            "安全漏洞名称",
            "风险级别",
            std_label,
            "安全漏洞代码位置",
            "安全漏洞所在位置截图证明",
        ]

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        header_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 设置表头背景色
        for cell in header_row.cells:
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "D9E2F3",
            })
            shading.append(shading_elm)

        # ── 数据行（按同源分组合并，同位置多个标准条款合并为一行） ──
        from collections import OrderedDict
        groups: dict[str, list] = OrderedDict()
        for f in confirmed:
            src = f.get("_source_group", f"{f.get('file_path','')}:{f.get('line_start',0)}")
            if src not in groups:
                groups[src] = []
            groups[src].append(f)

        for idx, (src_key, group_findings) in enumerate(groups.items(), 1):
            primary = group_findings[0]
            vuln_name = primary.get("vuln_name", "")

            # 合并条款号: "6.2.3.4/8.3.2"
            clauses = list(OrderedDict.fromkeys(
                f.get("clause", "") for f in group_findings
            ))
            clause_str = "/".join(clauses) + " " + vuln_name if clauses else vuln_name

            # 风险级别取组中最高
            severity = "低"
            for f in group_findings:
                s = f.get("severity", "低")
                if s == "高": severity = "高"; break
                if s == "中": severity = "中"

            # 标准类型列（合并条款号: "6.2.3.4/8.3.2 SQL注入"）
            clauses = list(OrderedDict.fromkeys(
                f.get("clause", "") for f in group_findings
            ))
            gb_type_str = "/".join(clauses) + " " + vuln_name if clauses else vuln_name

            row = table.add_row()
            self._set_cell(row.cells[0], str(idx), align="center")
            self._set_cell(row.cells[1], vuln_name)
            self._set_cell(row.cells[2], severity, align="center")
            self._set_cell(row.cells[3], gb_type_str)
            self._set_cell(row.cells[4], self._loc_str(primary))
            self._set_cell(row.cells[5], "")

        # ── 列宽设置 ──
        widths = [Cm(1.0), Cm(3.5), Cm(1.5), Cm(4.5), Cm(4.5), Cm(3.0)]
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w

        # ── 结束语 ──
        doc.add_paragraph()  # 空行
        end = doc.add_paragraph()
        end.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = end.add_run("…………………………………本结果报告单结束……………………………………")
        run.font.size = Pt(10)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        doc.save(str(path))

    @staticmethod
    def _standard_header(result: Any) -> str:
        """根据 result.standard 动态生成 DOCX 表头第 4 列。

        --standard 39412  → "对应GB/T 39412标准漏洞类型"
        --standard 34943  → "对应GB/T 34943标准漏洞类型"
        未指定（默认 ""） → "对应GB/T 34944标准漏洞类型"
        """
        std_map = {
            "39412": "GB/T 39412",
            "34944": "GB/T 34944",
            "34943": "GB/T 34943",
        }
        std = getattr(result, "standard", "") or ""
        name = std_map.get(std, "GB/T 34944")
        return f"对应{name}标准漏洞类型"

    @staticmethod
    def _loc_str(finding: dict) -> str:
        """从 finding 提取 CNAS 格式的代码位置字符串。

        taint mode 发现：显示入口点 + 爆发点
          格式: "（1）入口点：{文件}第{行}行\n（2）爆发点：{文件}第{行}行"

        普通发现：仅显示爆发点
          格式: "{文件名}第{行}行"
        """
        entry = finding.get("entry_point", {})
        outbreak_file = os.path.basename(finding.get("file_path", ""))
        outbreak_line = finding.get("line_start", 0)

        if entry and entry.get("file") and entry.get("line"):
            entry_file = os.path.basename(entry["file"])
            entry_line = entry["line"]
            return (
                f"（1）入口点：{entry_file}第{entry_line}行\n"
                f"（2）爆发点：{outbreak_file}第{outbreak_line}行"
            )

        return f"{outbreak_file}第{outbreak_line}行"

    @staticmethod
    def _set_cell(cell, text: str, align: str = "left") -> None:
        """设置表格单元格文本和对齐方式。"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        cell.text = ""
        p = cell.paragraphs[0]
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ─── HTML 报告（可选）─────────────────────────────────────────

    def _generate_html(self, result: Any, path: Path) -> None:
        """生成简单的 HTML 报告。"""
        md_lines: list[str] = []
        self._md_h1(md_lines, "CNAS 源代码安全审计报告")
        self._md_section(md_lines, "一、审计概要", self._render_summary_table(result, result.metadata if hasattr(result, "metadata") else None))
        self._md_section(md_lines, "二、漏洞发现汇总", self._render_findings_summary(result.findings if hasattr(result, "findings") else []))
        self._md_section(md_lines, "三、漏洞详情", self._render_findings_detail(result.findings if hasattr(result, "findings") else []))

        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="UTF-8"><title>CNAS 源代码安全审计报告</title>
<style>
  body {{ font-family: -apple-system, "Microsoft YaHei", sans-serif; max-width: 960px; margin: 0 auto; padding: 20px; color: #333; }}
  h1 {{ border-bottom: 2px solid #1a56db; padding-bottom: 10px; }}
  h2 {{ color: #1a56db; margin-top: 30px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
  th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
  th {{ background: #f5f7fa; }}
  .severity-high {{ color: #dc2626; font-weight: bold; }}
  .severity-medium {{ color: #d97706; font-weight: bold; }}
  .severity-low {{ color: #6b7280; }}
  pre {{ background: #f5f7fa; padding: 12px; border-radius: 4px; overflow-x: auto; }}
</style></head>
<body>
{"".join(f"<p>{line}</p>" if line and not line.startswith("#") and not line.startswith("|") and not line.startswith("-") and not line.startswith("  ") else f"<h1>{line[2:]}</h1>" if line.startswith("# ") else f"<h2>{line[3:]}</h2>" if line.startswith("## ") else line for line in md_lines)}
</body></html>"""

        with open(path, "w", encoding="utf-8") as f:
            f.write(html)

    # ─── Markdown 构建辅助 ────────────────────────────────────────

    @staticmethod
    def _md_h1(lines: list[str], text: str) -> None:
        lines.append(f"# {text}")
        lines.append("")

    @staticmethod
    def _md_h2(lines: list[str], text: str) -> None:
        lines.append(f"## {text}")
        lines.append("")

    @staticmethod
    def _md_section(lines: list[str], title: str, content: str) -> None:
        if content.strip():
            ReportGenerator._md_h2(lines, title)
            lines.append(content)
            lines.append("")

    # ─── 审计概要 ─────────────────────────────────────────────────

    def _render_summary_table(self, result: Any, metadata: Any) -> str:
        """渲染审计概要表格。"""
        mode = getattr(result, "mode", "unknown")
        duration = getattr(result, "duration_seconds", 0.0)
        findings = result.findings if hasattr(result, "findings") else []
        confirmed = [f for f in findings if f.get("auto_confidence", 0) >= 0.7]
        suspects = [f for f in findings if 0.4 <= f.get("auto_confidence", 0) < 0.7]

        llm_note = "在线模式（LLM 增强）" if mode == "online" else "离线模式（纯 SAST）"

        lines = [
            "| 项目 | 内容 |",
            "|------|------|",
            f"| 审计时间 | {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} |",
            f"| 运行模式 | {llm_note} |",
            f"| 审计耗时 | {duration:.1f} 秒 |",
        ]

        if metadata:
            lines.append(f"| 语言分布 | {', '.join(metadata.languages_detected) if hasattr(metadata, 'languages_detected') else 'N/A'} |")
            lines.append(f"| 文件总数 | {getattr(metadata, 'total_files', 0)} |")
            lines.append(f"| 构建系统 | {getattr(metadata, 'build_system', '无') or '无'} |")

        lines.append(f"| 确认漏洞 | {len(confirmed)} 条 |")
        lines.append(f"| 疑似漏洞 | {len(suspects)} 条 |")
        lines.append(f"| 发现总数 | {len(findings)} 条 |")

        if hasattr(result, "warnings") and result.warnings:
            lines.append(f"| 警告 | {len(result.warnings)} 条 |")

        lines.append("")
        return "\n".join(lines)

    # ─── 覆盖摘要 ─────────────────────────────────────────────────

    def _render_coverage_summary(self, findings: list[dict]) -> str:
        """渲染覆盖矩阵摘要。"""
        if not findings:
            return "_本次扫描未发现任何漏洞。_"

        # 按工具统计
        tool_counts: dict[str, int] = {}
        clause_set: set[str] = set()
        for f in findings:
            tool = f.get("source_tool", "unknown")
            tool_counts[tool] = tool_counts.get(tool, 0) + 1
            clause_set.add(f.get("clause", ""))

        lines = [
            "| 工具 | 发现数 |",
            "|------|--------|",
        ]
        for tool, count in sorted(tool_counts.items()):
            lines.append(f"| {tool} | {count} |")
        lines.append(f"| **合计** | **{len(findings)}** |")
        lines.append("")
        lines.append(f"覆盖条款数: {len(clause_set)}")
        lines.append("")
        return "\n".join(lines)

    # ─── 漏洞汇总 ─────────────────────────────────────────────────

    def _render_findings_summary(self, findings: list[dict]) -> str:
        """渲染漏洞按严重等级 + 按类别的汇总表。"""
        if not findings:
            return "_本次扫描未发现任何漏洞。_"

        # 按严重等级统计
        sev_counts: dict[str, int] = {"高": 0, "中": 0, "低": 0}
        for f in findings:
            sev = f.get("severity", "中")
            sev_counts[sev] = sev_counts.get(sev, 0) + 1

        # 按大类统计
        cat_counts: dict[str, int] = {}
        for f in findings:
            cat = f.get("category", "未分类")
            cat_counts[cat] = cat_counts.get(cat, 0) + 1

        lines = [
            "### 按严重等级统计",
            "",
            "| 严重等级 | 数量 |",
            "|---------|------|",
            f"| 🔴 高 | {sev_counts.get('高', 0)} |",
            f"| 🟡 中 | {sev_counts.get('中', 0)} |",
            f"| 🟢 低 | {sev_counts.get('低', 0)} |",
            f"| **合计** | **{len(findings)}** |",
            "",
            "### 按漏洞大类统计",
            "",
            "| 大类 | 数量 |",
            "|------|------|",
        ]
        for cat, count in sorted(cat_counts.items(), key=lambda x: -x[1]):
            lines.append(f"| {cat} | {count} |")
        lines.append("")
        return "\n".join(lines)

    # ─── 漏洞详情 ─────────────────────────────────────────────────

    def _render_findings_detail(self, findings: list[dict]) -> str:
        """渲染每条漏洞的详细信息，同源发现合并显示。"""
        if not findings:
            return "_本次扫描未发现任何漏洞。_"

        confirmed = [f for f in findings if f.get("auto_confidence", 0) >= 0.7]
        if not confirmed:
            return "_无高置信度的确认漏洞。_"

        # 按 _source_group 分组（同源发现的多个标准条款共享同一代码位置）
        grouped: dict[str, list[dict]] = {}
        ungrouped: list[dict] = []
        for f in confirmed:
            src = f.get("_source_group", "")
            if src:
                grouped.setdefault(src, []).append(f)
            else:
                ungrouped.append(f)

        lines: list[str] = []
        idx = 0

        # 渲染分组发现（跨标准同源）
        for src_group, group in grouped.items():
            idx += 1
            primary = group[0]
            clause = primary.get("clause", "")
            name = primary.get("vuln_name", "未知")
            severity = primary.get("severity", "中")
            sev_icon = {"高": "🔴", "中": "🟡", "低": "🟢"}.get(severity, "")
            confidence = primary.get("auto_confidence", 0.0)
            file_path = primary.get("file_path", "")
            line_start = primary.get("line_start", 0)
            line_end = primary.get("line_end", 0)
            source_tool = primary.get("source_tool", "")
            code = primary.get("code_snippet", "")

            lines.append("---")
            lines.append(f"### [{idx}] {sev_icon} {name}（同源，共 {len(group)} 个标准引用）")
            lines.append("")

            # 标准条款号表格
            lines.append("| 标准 | 条款号 |")
            lines.append("|------|--------|")
            for g in group:
                lines.append(f"| {g.get('standard', '')} | {g.get('clause', '')} |")
            lines.append("")

            lines.append(f"- **严重等级**: {severity}")
            lines.append(f"- **置信度**: {confidence:.0%}")
            lines.append(f"- **文件**: `{file_path}` (第 {line_start}-{line_end} 行)")
            lines.append(f"- **工具来源**: {source_tool}")

            if code:
                lines.append(f"```java")
                lines.append(code.strip())
                lines.append(f"```")

            # 知识库补充
            if self._kb and clause:
                vuln = self._kb.get_by_clause(clause)
                if vuln:
                    risk = vuln.get("risk", "")
                    fix = vuln.get("fix", "")
                    if risk:
                        lines.append(f"**风险描述**: {risk}")
                    if fix:
                        lines.append(f"**修复建议**: {fix[:200]}")
            lines.append("")

        # 渲染非分组发现
        for f in ungrouped:
            idx += 1
            lines.append("---")
            lines.append(self._render_single_finding_detail(idx, f))

        return "\n".join(lines)

    def _render_single_finding_detail(self, idx: int, f: dict) -> str:
        """渲染单条发现详情。"""
        clause = f.get("clause", "")
        name = f.get("vuln_name", "未知")
        severity = f.get("severity", "中")
        sev_icon = {"高": "🔴", "中": "🟡", "低": "🟢"}.get(severity, "")
        confidence = f.get("auto_confidence", 0.0)
        file_path = f.get("file_path", "")
        line_start = f.get("line_start", 0)
        line_end = f.get("line_end", 0)
        source_tool = f.get("source_tool", "")
        code = f.get("code_snippet", "")
        standard = f.get("standard", "")

        llm_info = ""
        if f.get("llm_status"):
            llm_info = f" | LLM: {f['llm_status']}"
            if f.get("llm_reasoning"):
                llm_info += f" ({f['llm_reasoning'][:80]})"

        lines = [
            f"### [{idx}] {sev_icon} {clause} {name}",
            "",
            f"- **严重等级**: {severity}",
            f"- **置信度**: {confidence:.0%}",
            f"- **标准**: {standard}",
            f"- **文件**: `{file_path}` (第 {line_start}-{line_end} 行)",
            f"- **工具来源**: {source_tool}{llm_info}",
        ]

        if code:
            lines.append("```java")
            lines.append(code.strip())
            lines.append("```")

        if self._kb and clause:
            vuln = self._kb.get_by_clause(clause)
            if vuln:
                risk = vuln.get("risk", "")
                fix = vuln.get("fix", "")
                if risk:
                    lines.append(f"**风险描述**: {risk}")
                if fix:
                    lines.append(f"**修复建议**: {fix[:200]}")
        lines.append("")
        return "\n".join(lines)

    # ─── 疑似发现 ─────────────────────────────────────────────────

    def _render_suspect_findings(self, findings: list[dict]) -> str:
        """渲染低置信度的疑似发现。"""
        suspects = [f for f in findings if 0.4 <= f.get("auto_confidence", 0) < 0.7]
        if not suspects:
            return "_无低置信度的疑似发现。_"

        lines = [f"_以下 {len(suspects)} 条发现置信度较低，建议人工审查:_", ""]
        for f in suspects:
            lines.append(
                f"- **{f.get('clause', '')} {f.get('vuln_name', '')}** "
                f"— `{f.get('file_path', '')}:{f.get('line_start', '')}` "
                f"(置信度: {f.get('auto_confidence', 0):.0%})"
            )
        lines.append("")
        return "\n".join(lines)

    # ─── 盲区 ─────────────────────────────────────────────────────

    def _render_blind_spots(self, result: Any) -> str:
        """渲染工具覆盖盲区和局限性说明。"""
        mode = getattr(result, "mode", "online")
        lines = [
            "### 方法局限性",
            "",
        ]

        if mode == "offline":
            lines.append("- ⚠️ **离线模式**：LLM 增强未执行，以下 9 种业务逻辑型漏洞无法检测：")
            lines.append("  - 6.2.5.2 违反信任边界、6.2.6.4 依赖Referer鉴权、6.2.6.11 反向域名解析")
            lines.append("  - 6.2.6.12 关键参数篡改、6.2.6.13 强口令要求")
            lines.append("  - 6.2.6.15 未验证Cookie、6.2.6.16 SQL关键字绕过授权")
            lines.append("  - 7.2.7.9/7.2.7.10 C/C++ 对应项")
            lines.append("")

        lines.append("- **SAST 工具固有局限**：无法检测需要业务上下文理解的逻辑缺陷")
        lines.append("- **CodeQL C/C++ 局限**：若编译失败将降级为纯语法分析，跨文件数据流精度下降")
        lines.append("")

        return "\n".join(lines)

    # ─── 元数据 ───────────────────────────────────────────────────

    def _render_metadata(self, result: Any, metadata: Any) -> str:
        """渲染审计元数据。"""
        duration = getattr(result, "duration_seconds", 0.0)
        mode = getattr(result, "mode", "unknown")

        lines = [
            "| 项目 | 内容 |",
            "|------|------|",
            f"| 运行模式 | {mode} |",
            f"| 总耗时 | {duration:.1f} 秒 |",
            f"| Run ID | {getattr(result, 'run_id', 'N/A')} |",
        ]

        if hasattr(result, "error_count"):
            lines.append(f"| 错误数 | {result.error_count} |")

        lines.append("")
        return "\n".join(lines)

    # ─── 辅助 ─────────────────────────────────────────────────────

    @staticmethod
    def _serialize_metadata(result: Any) -> dict[str, Any]:
        """将 ScanMetadata 序列化为字典。"""
        meta = result.metadata if hasattr(result, "metadata") else None
        if meta is None:
            return {}

        return {
            "languages_detected": getattr(meta, "languages_detected", []),
            "total_files": getattr(meta, "total_files", 0),
            "java_files": getattr(meta, "java_files", 0),
            "cpp_files": getattr(meta, "cpp_files", 0),
            "build_system": getattr(meta, "build_system", ""),
            "compile_ready": getattr(meta, "compile_ready", False),
            "source_roots": getattr(meta, "source_roots", []),
        }

    @staticmethod
    def _build_summary(result: Any) -> dict[str, Any]:
        """构建发现汇总统计。"""
        findings = result.findings if hasattr(result, "findings") else []
        total = len(findings)
        confirmed = sum(1 for f in findings if f.get("auto_confidence", 0) >= 0.7)
        suspects = sum(1 for f in findings if 0.4 <= f.get("auto_confidence", 0) < 0.7)

        sev_counts = {"高": 0, "中": 0, "低": 0}
        for f in findings:
            sev = f.get("severity", "中")
            sev_counts[sev] = sev_counts.get(sev, 0) + 1

        tool_counts: dict[str, int] = {}
        for f in findings:
            tool = f.get("source_tool", "unknown")
            tool_counts[tool] = tool_counts.get(tool, 0) + 1

        return {
            "total": total,
            "confirmed": confirmed,
            "suspects": suspects,
            "by_severity": sev_counts,
            "by_tool": tool_counts,
        }
